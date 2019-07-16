# -*- coding: utf-8 -*-
import xlwt
import xlrd
from xlutils.copy import copy
import pandas as pd
import numpy as np
import pymysql
from docx import Document
import logging
import pandas.io.sql as sql
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
from datetime import datetime,date
import time
import math
import re
import requests
from lxml import etree
import shutil
import os
import win32com.client
from aes_encrypt import aes_encrypt
logging.basicConfig(level=logging.DEBUG,
                    format='%(asctime)s %(filename)s[line:%(lineno)d] %(levelname)s %(message)s',
                    datefmt='%a, %d %b %Y %H:%M:%S',
                    filename='D:/uploadTempFiles/python/logFile/toWord.log',
                    filemode='a')


#Excel模板中的SQL
def createSqlStr_withParam(ent_tax_id,dateTm):
    sqlStrList = []
    #该行业内企业平均销售额增长率
    sqlStr1 = " SELECT amt1 "\
            +"    , round((amt1 / cnt1 - amt2 / cnt2) / (amt2 / cnt2), 4) AS growth_rate "\
            +"FROM ( "\
            +"    SELECT SUM(amount) AS amt1, COUNT(DISTINCT fp.ent_tax_id) AS cnt1 "\
            +"    FROM l3_receipt_sum fp "\
            +"        JOIN dim_ent_info ent ON fp.ent_tax_id = ent.ent_tax_id "\
            +"        JOIN ( "\
            +"            SELECT indu_cd "\
            +"            FROM dim_ent_info "\
            +"            WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +"        ) indu "\
            +"        ON ent.indu_cd = indu.indu_cd "\
            +"    WHERE receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7) "\
            +"        AND receipt_mon < LEFT('" + dateTm + "', 7) "\
            +") t1 "\
            +"    LEFT JOIN ( "\
            +"        SELECT SUM(amount) AS amt2, COUNT(DISTINCT fp.ent_tax_id) AS cnt2 "\
            +"        FROM l3_receipt_sum fp "\
            +"            JOIN dim_ent_info ent ON fp.ent_tax_id = ent.ent_tax_id "\
            +"            JOIN ( "\
            +"                SELECT indu_cd "\
            +"                FROM dim_ent_info "\
            +"                WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +"            ) indu "\
            +"            ON ent.indu_cd = indu.indu_cd "\
            +"        WHERE receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -2 YEAR), 7) "\
            +"            AND receipt_mon < LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7) "\
            +"    ) t2 "\
            +"    ON 1 = 1; "
    sqlStrList.append(sqlStr1)
    # 销售额本地行业排名
    sqlStr2 = " SELECT round(indu.indu_ent_cnt * (SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / COUNT(*), 0) AS rank1 "\
            +"  , round((SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / COUNT(*), 4) AS pct "\
            +" FROM ( "\
            +"  SELECT fp.ent_tax_id, SUM(amount) AS ind_value, min(indu.indu_ent_cnt) as indu_ent_cnt "\
            +"  FROM l3_receipt_sum fp "\
            +"      JOIN dim_ent_info ent ON fp.ent_tax_id = ent.ent_tax_id "\
            +"      JOIN ( "\
            +"          SELECT indu_cd, indu_ent_cnt "\
            +"          FROM dim_ent_info "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +"      ) indu "\
            +"      ON ent.indu_cd = indu.indu_cd "\
            +"  WHERE receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"      AND receipt_mon < left('" + dateTm + "',7) "\
            +"  GROUP BY ent_tax_id "\
            +"  ORDER BY ind_value DESC "\
            +" ) indu "\
            +"  JOIN ( "\
            +"      SELECT SUM(amount) AS ind_value "\
            +"      FROM l3_receipt_sum fp "\
            +"      WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"          AND receipt_mon < left('" + dateTm + "',7) "\
            +"          AND ent_tax_id = '" + ent_tax_id + "') "\
            +"  ) target "\
            +"  ON 1 = 1; "
    sqlStrList.append(sqlStr2)
    #上游企业数量本地行业排名
    sqlStr3 = " SELECT CASE "\
            +"        WHEN fp.ind_value IS NULL THEN ent.indu_ent_cnt "\
            +"        ELSE round(ent.indu_ent_cnt * (SUM(CASE "\
            +"            WHEN indu.ind_value > fp.ind_value THEN 1 "\
            +"            ELSE 0 "\
            +"        END) + 1) / COUNT(*), 0) "\
            +"    END AS rank1 "\
            +"    , CASE  "\
            +"        WHEN fp.ind_value IS NULL THEN 1.0 "\
            +"        ELSE round((SUM(CASE  "\
            +"            WHEN indu.ind_value > fp.ind_value THEN 1 "\
            +"            ELSE 0 "\
            +"        END) + 1) / COUNT(*), 4) "\
            +"    END AS pct "\
            +"FROM ( "\
            +"    SELECT ent_tax_id, indu_ent_cnt "\
            +"    FROM dim_ent_info "\
            +"    WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +") ent "\
            +"    LEFT JOIN ( "\
            +"        SELECT ent_tax_id, coalesce(COUNT(DISTINCT sale_ent_id), 0) AS ind_value "\
            +"        FROM l1_purchase_detail fp "\
            +"        WHERE (deduct_period >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7) "\
            +"            AND deduct_period < LEFT('" + dateTm + "', 7) "\
            +"            and deduct_flg != '*' "\
            +"            AND ent_tax_id = '" + ent_tax_id + "') "\
            +"    ) fp "\
            +"    ON ent.ent_tax_id = fp.ent_tax_id "\
            +"    LEFT JOIN ( "\
            +"        SELECT fp.ent_tax_id, coalesce(COUNT(DISTINCT sale_ent_id), 0) AS ind_value "\
            +"            , MIN(indu.indu_ent_cnt) AS indu_ent_cnt "\
            +"        FROM l1_purchase_detail fp "\
            +"            JOIN dim_ent_info ent ON fp.ent_tax_id = ent.ent_tax_id "\
            +"            JOIN ( "\
            +"                SELECT indu_cd, indu_ent_cnt "\
            +"                FROM dim_ent_info "\
            +"                WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +"            ) indu "\
            +"            ON ent.indu_cd = indu.indu_cd "\
            +"        WHERE deduct_period >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7) "\
            +"            AND deduct_period < LEFT('" + dateTm + "', 7) "\
            +"            and deduct_flg != '*' "\
            +"        GROUP BY fp.ent_tax_id "\
            +"        ORDER BY ind_value DESC "\
            +"    ) indu "\
            +"    ON 1 = 1; "
    sqlStrList.append(sqlStr3)
    #下游企业数量本地行业排名
    sqlStr4 = " SELECT round(indu.indu_ent_cnt * (SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / COUNT(*), 0) AS rank1 "\
            +"  , round((SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / COUNT(*), 4) AS pct "\
            +" FROM ( "\
            +"  SELECT fp.ent_tax_id, count(distinct buy_ent_nm) AS ind_value "\
            +"      , min(indu.indu_ent_cnt) as indu_ent_cnt "\
            +"  FROM l3_receipt_sum fp "\
            +"      JOIN dim_ent_info ent ON fp.ent_tax_id = ent.ent_tax_id "\
            +"      JOIN ( "\
            +"          SELECT indu_cd, indu_ent_cnt "\
            +"          FROM dim_ent_info "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +"      ) indu "\
            +"      ON ent.indu_cd = indu.indu_cd "\
            +"  WHERE receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"      AND receipt_mon < left('" + dateTm + "',7) "\
            +"  GROUP BY ent_tax_id "\
            +"  ORDER BY ind_value DESC "\
            +" ) indu "\
            +"  JOIN ( "\
            +"      SELECT count(distinct buy_ent_nm) AS ind_value "\
            +"      FROM l3_receipt_sum fp "\
            +"      WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"          AND receipt_mon < left('" + dateTm + "',7) "\
            +"          AND ent_tax_id = '" + ent_tax_id + "') "\
            +"  ) target "\
            +"  ON 1 = 1; "
    sqlStrList.append(sqlStr4)
    #近一年增值税纳税额本地行业地位
    sqlStr5 = " SELECT round(indu.indu_ent_cnt * (SUM(CASE                                                                                              "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1                                                        "\
            +"      ELSE 0                                                                                               "\
            +"  END) + 1) / COUNT(*), 0) AS rank1                                                                                         "\
            +"  , case when target.ind_value=0 then 1.0 else               "\
            +"  round((SUM(CASE                                                                                        "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1                                                        "\
            +"      ELSE 0                                                                                               "\
            +"  END) + 1) / COUNT(*), 4) end AS pct                                                                          "\
            +" FROM (                                                                                                       "\
            +"  SELECT fp.ent_tax_id, SUM(tax_amt) AS ind_value                                                          "\
            +"      , min(indu.indu_ent_cnt) as indu_ent_cnt                   "\
            +"  FROM l1_all_list fp                                                                                      "\
            +"      JOIN dim_ent_info ent ON fp.ent_tax_id = ent.ent_tax_id                                              "\
            +"      JOIN (                                                                                               "\
            +"          SELECT indu_cd, indu_ent_cnt                                                                                   "\
            +"          FROM dim_ent_info                                                                                "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "'                                                          "\
            +"      ) indu                                                                                               "\
            +"      ON ent.indu_cd = indu.indu_cd                                                                        "\
            +"  WHERE end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                                             "\
            +"      AND end_date < '" + dateTm + "'                                                                          "\
            +"         and item_nm like '%增值税%'                                                                     "\
            +"  GROUP BY ent_tax_id                                                                                      "\
            +"  ORDER BY ind_value DESC                                                                                  "\
            +" ) indu                                                                                                       "\
            +"  JOIN (                                                                                                   "\
            +"      SELECT SUM(tax_amt) AS ind_value                                                                     "\
            +"      FROM l1_all_list fp                                                                                  "\
            +"      WHERE (end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                                        "\
            +"          AND end_date < '" + dateTm + "'                                                                      "\
            +"             and item_nm like '%增值税%'                                                                 "\
            +"          AND ent_tax_id = '" + ent_tax_id + "')                                                           "\
            +"  ) target                                                                                                 "\
            +"  ON 1 = 1;                                                                                                "
    sqlStrList.append(sqlStr5)
    #近一年增值税税负率本地行业地位
    sqlStr6 = " SELECT round(indu.indu_ent_cnt * (SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / SUM(case when indu.ind_value>0 then 1 else 0 end), 0) AS rank1 "\
            +"  , case when target.ind_value=0 then 1.0 else "\
            +"      round((SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / SUM(case when indu.ind_value>0 then 1 else 0 end), 4) end AS pct "\
            +" FROM ( "\
            +"  SELECT sale.ent_tax_id, coalesce(tax.ind_value, 0) / sale.ind_value AS ind_value "\
            +"      , indu.indu_ent_cnt as indu_ent_cnt "\
            +"  FROM ( "\
            +"      SELECT fp.ent_tax_id, SUM(amount) AS ind_value "\
            +"      FROM l3_receipt_sum fp "\
            +"      WHERE receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"          AND receipt_mon < left('" + dateTm + "',7) "\
            +"      GROUP BY ent_tax_id "\
            +"  ) sale "\
            +"      LEFT JOIN ( "\
            +"          SELECT fp.ent_tax_id, SUM(tax_amt) AS ind_value "\
            +"          FROM l1_all_list fp "\
            +"          WHERE (end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR) "\
            +"              AND end_date < '" + dateTm + "' "\
            +"              AND item_nm LIKE '%增值税%') "\
            +"          GROUP BY ent_tax_id "\
            +"      ) tax "\
            +"      ON sale.ent_tax_id = tax.ent_tax_id "\
            +"      JOIN dim_ent_info ent ON sale.ent_tax_id = ent.ent_tax_id "\
            +"      JOIN ( "\
            +"          SELECT indu_cd, indu_ent_cnt "\
            +"          FROM dim_ent_info "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +"      ) indu "\
            +"      ON ent.indu_cd = indu.indu_cd "\
            +"  ORDER BY ind_value DESC "\
            +" ) indu "\
            +"  JOIN ( "\
            +"      SELECT sale.ent_tax_id, coalesce(tax.ind_value, 0) / sale.ind_value AS ind_value "\
            +"      FROM ( "\
            +"          SELECT fp.ent_tax_id, SUM(amount) AS ind_value "\
            +"          FROM l3_receipt_sum fp "\
            +"          WHERE receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"              AND receipt_mon < left('" + dateTm + "',7) "\
            +"              AND ent_tax_id = '" + ent_tax_id + "' "\
            +"          GROUP BY ent_tax_id "\
            +"      ) sale "\
            +"          LEFT JOIN ( "\
            +"              SELECT fp.ent_tax_id, SUM(tax_amt) AS ind_value "\
            +"              FROM l1_all_list fp "\
            +"              WHERE (end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR) "\
            +"                  AND end_date < '" + dateTm + "' "\
            +"                  AND item_nm LIKE '%增值税%' "\
            +"                  AND ent_tax_id = '" + ent_tax_id + "') "\
            +"              GROUP BY ent_tax_id "\
            +"          ) tax "\
            +"          ON sale.ent_tax_id = tax.ent_tax_id "\
            +"  ) target "\
            +"  ON 1 = 1; "
    sqlStrList.append(sqlStr6)
    #毛利率本地行业地位
    sqlStr7 = " SELECT round(indu.indu_ent_cnt * (SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value and indu.ind_value!=0 THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / SUM(case when indu.ind_value>0 then 1 else 0 end), 0) AS rank1 "\
            +"  , round((SUM(CASE "\
            +"      WHEN indu.ind_value > target.ind_value and indu.ind_value!=0 THEN 1 "\
            +"      ELSE 0 "\
            +"  END) + 1) / SUM(case when indu.ind_value!=0 then 1 else 0 end), 4) AS pct "\
            +" FROM ( "\
            +"  SELECT sale.ent_tax_id, "\
            +"     case when coalesce(tax.ind_value, 0)=0 then 0 else (sale.ind_value-tax.ind_value)/sale.ind_value end AS ind_value "\
            +"     , indu.indu_ent_cnt as indu_ent_cnt "\
            +"  FROM ( "\
            +"      SELECT fp.ent_tax_id, SUM(amount) AS ind_value "\
            +"      FROM l3_receipt_sum fp "\
            +"      WHERE receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"          AND receipt_mon < left('" + dateTm + "',7) "\
            +"      GROUP BY ent_tax_id "\
            +"  ) sale "\
            +"      LEFT JOIN ( "\
            +"          SELECT ent_tax_id, SUM(amount) AS ind_value "\
            +"          FROM rdc.l1_purchase_tax "\
            +"          WHERE (end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR) "\
            +"              AND end_date < '" + dateTm + "' "\
            +"              AND item_no = 12) "\
            +"          GROUP BY ent_tax_id "\
            +"      ) tax "\
            +"      ON sale.ent_tax_id = tax.ent_tax_id "\
            +"      JOIN dim_ent_info ent ON sale.ent_tax_id = ent.ent_tax_id "\
            +"      JOIN ( "\
            +"          SELECT indu_cd, indu_ent_cnt "\
            +"          FROM dim_ent_info "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "' "\
            +"      ) indu "\
            +"      ON ent.indu_cd = indu.indu_cd "\
            +"  ORDER BY ind_value DESC "\
            +" ) indu "\
            +"  JOIN ( "\
            +"      SELECT sale.ent_tax_id, "\
            +"         case when coalesce(tax.ind_value, 0)=0 then 0 else (sale.ind_value-tax.ind_value)/sale.ind_value end AS ind_value "\
            +"      FROM ( "\
            +"          SELECT fp.ent_tax_id, SUM(amount) AS ind_value "\
            +"          FROM l3_receipt_sum fp "\
            +"          WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"              AND receipt_mon < left('" + dateTm + "',7) "\
            +"              AND ent_tax_id = '" + ent_tax_id + "') "\
            +"          GROUP BY ent_tax_id "\
            +"      ) sale "\
            +"          LEFT JOIN ( "\
            +"              SELECT ent_tax_id, SUM(amount) AS ind_value "\
            +"              FROM rdc.l1_purchase_tax "\
            +"              WHERE (end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR) "\
            +"                  AND end_date < '" + dateTm + "' "\
            +"                  AND item_no = 12 "\
            +"                  AND ent_tax_id = '" + ent_tax_id + "') "\
            +"          ) tax "\
            +"          ON sale.ent_tax_id = tax.ent_tax_id "\
            +"  ) target "\
            +"  ON 1 = 1; "
    sqlStrList.append(sqlStr7)
    #近一年最大十家稳定上游供应商数量
    sqlStr8 = " SELECT '', COUNT(*) CNT                                                                                                  "\
            +" FROM (                                                                                                    "\
            +"  SELECT c.sale_ent_id                                                                                     "\
            +"  FROM (                                                                                                   "\
            +"      SELECT sale_ent_id, COUNT(DISTINCT monthsid) AS cnt                                                  "\
            +"      FROM (                                                                                               "\
            +"          SELECT DISTINCT sale_ent_id                                                                      "\
            +"              , floor((substring(receipt_date, 6, 2) + 0) / 3) AS monthsid                                 "\
            +"          FROM rdc.l1_purchase_detail                                                                      "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "'                                                          "\
            +"            AND deduct_period >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                     "\
            +"            AND deduct_period < LEFT('" + dateTm + "', 7)                                                  "\
            +"            AND deduct_flg != '*'                                                                          "\
            +"      ) a                                                                                                  "\
            +"      GROUP BY sale_ent_id                                                                                 "\
            +"  ) c                                                                                                      "\
            +"      JOIN (                                                                                               "\
            +"          SELECT DISTINCT COUNT(DISTINCT floor((substring(receipt_date, 6, 2) + 0) / 3)) AS cnt            "\
            +"          FROM rdc.l1_purchase_detail                                                                      "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "'                                                          "\
            +"            AND deduct_period >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                     "\
            +"            AND deduct_period < LEFT('" + dateTm + "', 7)                                                  "\
            +"            AND deduct_flg != '*'                                                                          "\
            +"      ) b                                                                                                  "\
            +"      ON c.cnt = b.cnt                                                                                     "\
            +"      JOIN (                                                                                               "\
            +"          SELECT sale_ent_id, SUM(amount) AS amt                                                           "\
            +"          FROM l1_purchase_detail                                                                          "\
            +"          WHERE ent_tax_id = '" + ent_tax_id + "'                                                          "\
            +"            AND deduct_period >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                     "\
            +"            AND deduct_period < LEFT('" + dateTm + "', 7)                                                  "\
            +"            AND deduct_flg != '*'                                                                          "\
            +"          GROUP BY sale_ent_id                                                                             "\
            +"          ORDER BY amt DESC                                                                                "\
            +"          LIMIT 10                                                                                         "\
            +"      ) d                                                                                                  "\
            +"      ON c.sale_ent_id = d.sale_ent_id                                                                     "\
            +" ) e;                                                                                                      "
    sqlStrList.append(sqlStr8)
    #近两年每年最大十家稳定下游交易对手数量
    sqlStr9 = " SELECT SUM(loc) AS loc, count(*) - COUNT(DISTINCT buy_ent_nm) AS cnt "\
            +"FROM ( "\
            +"  (SELECT buy_ent_nm, SUM(amount) AS ind_value,1 as loc "\
            +"  FROM l3_receipt_sum "\
            +"  WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"      AND receipt_mon < left('" + dateTm + "',7) "\
            +"      AND ent_tax_id = '" + ent_tax_id + "') "\
            +"  GROUP BY buy_ent_nm "\
            +"  ORDER BY ind_value DESC "\
            +"  LIMIT 10) "\
            +"  UNION ALL "\
            +"  (SELECT buy_ent_nm, SUM(amount) AS ind_value,100 as loc "\
            +"  FROM l3_receipt_sum "\
            +"  WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -2 YEAR),7) "\
            +"      AND receipt_mon < left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +"      AND ent_tax_id = '" + ent_tax_id + "') "\
            +"  GROUP BY buy_ent_nm "\
            +"  ORDER BY ind_value DESC "\
            +"  LIMIT 10) "\
            +") d; "
    sqlStrList.append(sqlStr9)
    #开票断档月份数
    sqlStr10 = " SELECT GROUP_CONCAT(dim_mon.mon SEPARATOR '，'), COUNT(*)                "\
            +" FROM dim_mon                                                               "\
            +"     LEFT JOIN (                                                            "\
            +"         SELECT DISTINCT receipt_mon AS mon                                 "\
            +"         FROM l3_receipt_sum                                                "\
            +"         WHERE ent_tax_id = '" + ent_tax_id + "'                            "\
            +"     ) detail                                                               "\
            +"     ON dim_mon.mon = detail.mon                                            "\
            +" WHERE (dim_mon.mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7) "\
            +" AND dim_mon.mon < left('" + dateTm + "',7)                                 "\
            +" AND detail.mon IS NULL);                                                   "
    sqlStrList.append(sqlStr10)
    # 销售额增长率
    sqlStr11 = " SELECT amt1 "\
            +"    , round((amt1 - amt2) / amt2, 4) AS growth_rate "\
            +"FROM ( "\
            +"    SELECT SUM(amount) AS amt1 "\
            +"    FROM l3_receipt_sum "\
            +"    WHERE (ent_tax_id = '" + ent_tax_id + "' "\
            +"        AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7) "\
            +"        AND receipt_mon < LEFT('" + dateTm + "', 7)) "\
            +") t1 "\
            +"    JOIN ( "\
            +"        SELECT SUM(amount) AS amt2 "\
            +"        FROM l3_receipt_sum "\
            +"        WHERE (ent_tax_id = '" + ent_tax_id + "' "\
            +"            AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -2 YEAR), 7) "\
            +"            AND receipt_mon < LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)) "\
            +"    ) t2 "\
            +"    ON 1 = 1; "
    sqlStrList.append(sqlStr11)
    # 主营商品增长率
    sqlStr12 = " SELECT amt1                                                                                             "\
            +"  , round((amt1 - amt2) / amt2, 4) AS growth_rate                                                          "\
            +" FROM (                                                                                                    "\
            +"  SELECT SUM(CASE                                                                                          "\
            +"          WHEN (receipt_flg IN ('0', '1')                                                                  "\
            +"          AND receipt_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                                 "\
            +"          AND receipt_date < '" + dateTm + "') THEN amount                                                 "\
            +"          ELSE 0                                                                                           "\
            +"      END) AS amt1, SUM(CASE                                                                               "\
            +"          WHEN (receipt_flg IN ('0', '1')                                                                  "\
            +"          AND receipt_date >= date_add('" + dateTm + "', INTERVAL -2 YEAR)                                 "\
            +"          AND receipt_date < date_add('" + dateTm + "', INTERVAL -1 YEAR)) THEN amount                     "\
            +"          ELSE 0                                                                                           "\
            +"      END) AS amt2                                                                                         "\
            +"  FROM l2_receipt_detail fp                                                                                "\
            +"      JOIN dim_ent_info ent                                                                                "\
            +"      ON fp.ent_tax_id = ent.ent_tax_id                                                                    "\
            +"          AND LEFT(fp.tax_code, 7) = LEFT(ent.pro_tax_cd, 7)                                               "\
            +"  WHERE fp.ent_tax_id = '" + ent_tax_id + "'                                                               "\
            +" ) a;                                                                                                      "
    sqlStrList.append(sqlStr12)
    # 增值税纳税额增长率
    sqlStr13 = " SELECT amt1                                                                                             "\
            +"  , round((amt1 - amt2) / amt2, 4) AS growth_rate                                                          "\
            +" FROM (                                                                                                    "\
            +"  SELECT SUM(CASE                                                                                          "\
            +"          WHEN end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                                    "\
            +"          AND end_date < '" + dateTm + "' THEN tax_amt                                                     "\
            +"          ELSE 0                                                                                           "\
            +"      END) AS amt1, SUM(CASE                                                                               "\
            +"          WHEN end_date >= date_add('" + dateTm + "', INTERVAL -2 YEAR)                                    "\
            +"          AND end_date < date_add('" + dateTm + "', INTERVAL -1 YEAR) THEN tax_amt                         "\
            +"          ELSE 0                                                                                           "\
            +"      END) AS amt2                                                                                         "\
            +"  FROM l1_all_list                                                                                         "\
            +"  WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                 "\
            +"      AND item_nm LIKE '%增值税%')                                                                         "\
            +" ) a                                                                                                       "
    sqlStrList.append(sqlStr13)

    return sqlStrList


def createSqlStr_for_word(ent_tax_id, dateTm):
    sqlStrList = []
    #0 #1 #53  所需参数：企业税号
    sqlStr0 = "select ent_nm, indu_nm, pro_nm from dim_ent_info where ent_tax_id = '" + ent_tax_id + "'"
    sqlStrList.append(sqlStr0)
    #9 #10  所需参数：企业税号、查询日期
    sqlStr1 = " select '供应商',count(distinct sale_ent_id) cnt ,round(sum(amount)/10000,2) amt                    "\
            +" from l1_purchase_detail                                                                             "\
            +" WHERE deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and ent_tax_id = '" + ent_tax_id + "' and deduct_flg != '*'                                    "
    sqlStrList.append(sqlStr1)
    #5 #16 #17
    sqlStr2 = " select '下游企业',count(distinct buy_ent_nm) cnt ,round(SUM(amount)/10000,2) amt                   "\
            +" from l3_receipt_sum                                                                                 "\
            +" WHERE receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                           "\
            +"      AND receipt_mon < left('" + dateTm + "',7)                                                     "\
            +"         and ent_tax_id = '" + ent_tax_id + "'                                                       "
    sqlStrList.append(sqlStr2)
    # 23  所需参数：企业税号、查询日期
    sqlStr3 = " select '增值税',count(distinct end_date) cnt,round(sum(tax_amt)/10000,2) amt                       "\
            +" from l1_all_list                                                                                    "\
            +"  WHERE end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                                     "\
            +"      AND end_date < '" + dateTm + "'                                                                "\
            +"      and item_nm like '%增值税%'                                                                    "\
            +"      and ent_tax_id = '" + ent_tax_id + "';                                                         "
    sqlStrList.append(sqlStr3)
    #32
    sqlStr4 = " SELECT sale.ent_tax_id,                                                                            "\
            +"case when coalesce(tax.ind_value, 0)=0 then 0 else (sale.ind_value-tax.ind_value)/sale.ind_value end AS ind_value    "\
            +" FROM (                                                                                              "\
            +"  SELECT fp.ent_tax_id, SUM(amount) AS ind_value                                                     "\
            +"  FROM l3_receipt_sum fp                                                                             "\
            +"  WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND receipt_mon < left('" + dateTm + "',7)                                                     "\
            +"         and ent_tax_id = '" + ent_tax_id + "')                                                      "\
            +"  GROUP BY ent_tax_id                                                                                "\
            +" ) sale                                                                                              "\
            +"  LEFT JOIN (                                                                                        "\
            +"      SELECT ent_tax_id, COUNT(*) AS cnt, SUM(amount) AS ind_value                                   "\
            +"      FROM rdc.l1_purchase_tax                                                                       "\
            +"      WHERE (end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                                "\
            +"          AND end_date < '" + dateTm + "'                                                            "\
            +"          AND item_no = 12                                                                           "\
            +"          AND ent_tax_id = '" + ent_tax_id + "')                                                     "\
            +"  ) tax                                                                                              "\
            +"  ON sale.ent_tax_id = tax.ent_tax_id;                                                               "

    sqlStrList.append(sqlStr4)
    #37
    sqlStr5 = " SELECT round(sum(max10.amt)/10000,2) as amt                                                 "\
            +" FROM (                                                                                       "\
            +"      SELECT sale_ent_id, SUM(amount) AS amt                                                  "\
            +"      FROM l1_purchase_detail                                                                 "\
            +"      WHERE deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)             "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                            "\
            +"      and ent_tax_id = '" + ent_tax_id + "' and deduct_flg != '*'                             "\
            +"      GROUP BY sale_ent_id                                                                    "\
            +"      ORDER BY amt DESC                                                                       "\
            +"      LIMIT 10                                                                                "\
            +"  ) max10                                                                                     "\
            +"  JOIN (                                                                                      "\
            +"      SELECT sale_ent_id, COUNT(DISTINCT monthsid) AS cnt                                     "\
            +"      FROM (                                                                                  "\
            +"          SELECT DISTINCT sale_ent_id                                                         "\
            +"              , floor((substring(receipt_date, 6, 2) + 0) / 3) AS monthsid                    "\
            +"          FROM l1_purchase_detail                                                             "\
            +"          WHERE deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)         "\
            +"          AND deduct_period < left('" + dateTm + "',7)                                        "\
            +"          and ent_tax_id = '" + ent_tax_id + "' and deduct_flg != '*'                         "\
            +"      ) a                                                                                     "\
            +"      GROUP BY sale_ent_id                                                                    "\
            +"  ) list                                                                                      "\
            +"  on max10.sale_ent_id = list.sale_ent_id                                                     "\
            +"  JOIN (                                                                                      "\
            +"      SELECT DISTINCT COUNT(DISTINCT floor((substring(receipt_date, 6, 2) + 0) / 3)) AS cnt   "\
            +"      FROM l1_purchase_detail                                                                 "\
            +"      WHERE deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                            "\
            +"      and ent_tax_id = '" + ent_tax_id + "' and deduct_flg != '*'                             "\
            +"  ) all_cnt                                                                                   "\
            +"  ON list.cnt = all_cnt.cnt;                                                                  "
    sqlStrList.append(sqlStr5)
    #42
    sqlStr6 = " SELECT round(sum(max10.ind_value) / 10000, 2) AS amt                                        "\
            +" FROM (                                                                                       "\
            +"  SELECT buy_ent_nm                                                                           "\
            +"  FROM (                                                                                      "\
            +"      (SELECT buy_ent_nm, SUM(amount) AS ind_value, 1 AS loc                                  "\
            +"      FROM l3_receipt_sum                                                                     "\
            +"      WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)              "\
            +"        AND receipt_mon < left('" + dateTm + "',7)                                            "\
            +"        and ent_tax_id = '" + ent_tax_id + "')                                                "\
            +"      GROUP BY buy_ent_nm                                                                     "\
            +"      ORDER BY ind_value DESC                                                                 "\
            +"      LIMIT 10)                                                                               "\
            +"      UNION ALL                                                                               "\
            +"      (SELECT buy_ent_nm, SUM(amount) AS ind_value, 100 AS loc                                "\
            +"      FROM l3_receipt_sum                                                                     "\
            +"      WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -2 YEAR),7)              "\
            +"        AND receipt_mon < left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                "\
            +"        AND ent_tax_id = '" + ent_tax_id + "')                                                "\
            +"      GROUP BY buy_ent_nm                                                                     "\
            +"      ORDER BY ind_value DESC                                                                 "\
            +"      LIMIT 10)                                                                               "\
            +"  ) d                                                                                         "\
            +"  GROUP BY buy_ent_nm                                                                         "\
            +"  HAVING COUNT(*) > 1                                                                         "\
            +" ) list                                                                                       "\
            +"  JOIN (                                                                                      "\
            +"      SELECT buy_ent_nm, SUM(amount) AS ind_value, 1 AS loc                                   "\
            +"      FROM l3_receipt_sum                                                                     "\
            +"      WHERE (receipt_mon >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)              "\
            +"        AND receipt_mon < left('" + dateTm + "',7)                                            "\
            +"        and ent_tax_id = '" + ent_tax_id + "')                                                "\
            +"      GROUP BY buy_ent_nm                                                                     "\
            +"      ORDER BY ind_value DESC                                                                 "\
            +"      LIMIT 10                                                                                "\
            +"  ) max10                                                                                     "\
            +"  ON list.buy_ent_nm = max10.buy_ent_nm;                                                      "
    sqlStrList.append(sqlStr6)
    #49 #51
    sqlStr7 = " SELECT round(amt1/10000,2) as amt1, round(amt2/10000,2) as amt2 "\
            +"FROM ( "\
            +"    SELECT SUM(amount) AS amt1 "\
            +"    FROM l3_receipt_sum "\
            +"    WHERE (ent_tax_id = '" + ent_tax_id + "' "\
            +"        AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7) "\
            +"        AND receipt_mon < LEFT('" + dateTm + "', 7)) "\
            +") t1 "\
            +"    JOIN ( "\
            +"        SELECT SUM(amount) AS amt2 "\
            +"        FROM l3_receipt_sum "\
            +"        WHERE (ent_tax_id = '" + ent_tax_id + "' "\
            +"            AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -2 YEAR), 7) "\
            +"            AND receipt_mon < LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)) "\
            +"    ) t2 "\
            +"    ON 1 = 1; "
    sqlStrList.append(sqlStr7)
    #56 #58
    sqlStr8 = " SELECT round(amt1/10000,2) as amt1                                                              "\
            +"  , round(amt2/10000,2) as amt2                                                               "\
            +" FROM (                                                                                          "\
            +"  SELECT SUM(CASE                                                                             "\
            +"          WHEN (receipt_flg IN ('0', '1')                                                     "\
            +"          AND receipt_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                         "\
            +"          AND receipt_date < '" + dateTm + "') THEN amount                                         "\
            +"          ELSE 0                                                                              "\
            +"      END) AS amt1, SUM(CASE                                                                  "\
            +"          WHEN (receipt_flg IN ('0', '1')                                                     "\
            +"          AND receipt_date >= date_add('" + dateTm + "', INTERVAL -2 YEAR)                         "\
            +"          AND receipt_date < date_add('" + dateTm + "', INTERVAL -1 YEAR)) THEN amount             "\
            +"          ELSE 0                                                                              "\
            +"      END) AS amt2                                                                            "\
            +"  FROM l2_receipt_detail fp                                                                   "\
            +"      JOIN dim_ent_info ent                                                                   "\
            +"      ON fp.ent_tax_id = ent.ent_tax_id                                                       "\
            +"          AND LEFT(fp.tax_code, 7) = LEFT(ent.pro_tax_cd, 7)                                  "\
            +"  WHERE fp.ent_tax_id = '" + ent_tax_id + "'                                                           "\
            +" ) a;                                                                                            "
    sqlStrList.append(sqlStr8)
    #62 #64
    sqlStr9 = " SELECT round(amt1/10000,2) as amt1                                                              "\
            +"  , round(amt2/10000,2) as amt2                                                               "\
            +" FROM (                                                                                          "\
            +"  SELECT SUM(CASE                                                                             "\
            +"          WHEN end_date >= date_add('" + dateTm + "', INTERVAL -1 YEAR)                          "\
            +"          AND end_date < '" + dateTm + "' THEN tax_amt                                             "\
            +"          ELSE 0                                                                              "\
            +"      END) AS amt1, SUM(CASE                                                                  "\
            +"          WHEN end_date >= date_add('" + dateTm + "', INTERVAL -2 YEAR)                          "\
            +"          AND end_date < date_add('" + dateTm + "', INTERVAL -1 YEAR) THEN tax_amt                 "\
            +"          ELSE 0                                                                              "\
            +"      END) AS amt2                                                                            "\
            +"  FROM l1_all_list                                                                            "\
            +"  WHERE (ent_tax_id = '" + ent_tax_id + "'                                                             "\
            +"      AND item_nm LIKE '%增值税%')                                                       "\
            +" ) a;                                                                                            "
    sqlStrList.append(sqlStr9)
    return sqlStrList


def get_table_data(ent_tax_id, dateTm):
    table_list = []
    sqlStrList = []
    #第一个表格
    sqlStr1 = " (SELECT mx.sale_ent_id, round(mx.amt / 10000, 2) AS amt                                                  "\
            +"  , concat(round(100 * mx.amt / ttl.amt, 2), '%') AS pct                                               "\
            +" FROM (                                                                                                   "\
            +"  SELECT sale_ent_id, SUM(amount) AS amt                                                               "\
            +"  FROM l1_purchase_detail                                                                              "\
            +"  WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and deduct_flg != '*')                                                                   "\
            +"  GROUP BY sale_ent_id                                                                                 "\
            +"  ORDER BY amt DESC                                                                                    "\
            +"  LIMIT 10                                                                                             "\
            +" ) mx                                                                                                     "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS amt                                                                        "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                  "\
            +"          AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"          AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"          and deduct_flg != '*')                                                               "\
            +"  ) ttl                                                                                                "\
            +"  ON 1 = 1                                                                                             "\
            +" order by amt desc limit 10)                                                                              "\
            +" union all                                                                                                "\
            +" (select '合计',round(sum(mx.amt) / 10000, 2) AS amt                                                      "\
            +"  , concat(round(100 * sum(mx.amt) / ttl.amt, 2), '%') AS pct                                          "\
            +" FROM (                                                                                                   "\
            +"  SELECT sale_ent_id, SUM(amount) AS amt                                                               "\
            +"  FROM l1_purchase_detail                                                                              "\
            +"  WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and deduct_flg != '*')                                                                   "\
            +"  GROUP BY sale_ent_id                                                                                 "\
            +"  ORDER BY amt DESC                                                                                    "\
            +"  LIMIT 10                                                                                             "\
            +" ) mx                                                                                                     "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS amt                                                                        "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                  "\
            +"          AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"          AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"          and deduct_flg != '*')                                                               "\
            +"  ) ttl                                                                                                "\
            +"  ON 1 = 1 limit 1);                                                                                   "
    sqlStrList.append(sqlStr1)
    # 第二个表格
    sqlStr2 = " (SELECT mx.buy_ent_nm, round(mx.amt / 10000, 2) AS amt                                                   "\
            +"  , concat(round(100 * mx.amt / ttl.amt, 2), '%') AS pct                                               "\
            +" FROM (                                                                                                   "\
            +"  SELECT buy_ent_nm, SUM(amount) AS amt                                                                                      "\
            +"  FROM l3_receipt_sum                                                                               "\
            +"  WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                  "\
            +"  GROUP BY buy_ent_nm                                                                                  "\
            +"  ORDER BY amt DESC                                                                                    "\
            +"  LIMIT 10                                                                                             "\
            +" ) mx                                                                                                     "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS amt                                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                  "\
            +"          AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                                  "\
            +"          AND receipt_mon < LEFT('" + dateTm + "', 7))                                                              "\
            +"  ) ttl                                                                                                "\
            +"  ON 1 = 1                                                                                             "\
            +" order by amt desc limit 10)                                                                              "\
            +" union all                                                                                                "\
            +" (select '合计',round(sum(mx.amt) / 10000, 2) AS amt                                                      "\
            +"  , concat(round(100 * sum(mx.amt) / ttl.amt, 2), '%') AS pct                                          "\
            +" FROM (                                                                                                   "\
            +"  SELECT buy_ent_nm, SUM(amount) AS amt                                                                                      "\
            +"  FROM l3_receipt_sum                                                                               "\
            +"  WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                  "\
            +"  GROUP BY buy_ent_nm                                                                                  "\
            +"  ORDER BY amt DESC                                                                                    "\
            +"  LIMIT 10                                                                                             "\
            +" ) mx                                                                                                     "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS amt                                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                  "\
            +"          AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"          AND receipt_mon < LEFT('" + dateTm + "', 7))                                                              "\
            +"  ) ttl                                                                                                "\
            +"  ON 1 = 1 limit 1);                                                                                   "\
            +"                                                                                                          "
    sqlStrList.append(sqlStr2)
    # 第三个表格
    sqlStr3 = " (SELECT list.sale_ent_id,round(max10.amt/10000,2) as amt,concat(round(100*max10.amt/ttl.amt,2),'%') pct  "\
            +" FROM (                                                                                                   "\
            +"      SELECT sale_ent_id, SUM(amount) AS amt                                                           "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"          AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"          and deduct_flg != '*')                                                                   "\
            +"      GROUP BY sale_ent_id                                                                             "\
            +"      ORDER BY amt DESC                                                                                "\
            +"      LIMIT 10                                                                                         "\
            +"  ) max10                                                                                              "\
            +"  JOIN (                                                                                               "\
            +"      SELECT sale_ent_id, COUNT(DISTINCT monthsid) AS cnt                                              "\
            +"      FROM (                                                                                           "\
            +"          SELECT DISTINCT sale_ent_id                                                                  "\
            +"              , floor((substring(receipt_date, 6, 2) + 0) / 3) AS monthsid                             "\
            +"          FROM l1_purchase_detail                                                                      "\
            +"          WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"          AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"          AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"          and deduct_flg != '*')                                                               "\
            +"      ) a                                                                                              "\
            +"      GROUP BY sale_ent_id                                                                             "\
            +"  ) list                                                                                               "\
            +"  on max10.sale_ent_id = list.sale_ent_id                                                              "\
            +"  JOIN (                                                                                               "\
            +"      SELECT DISTINCT COUNT(DISTINCT floor((substring(receipt_date, 6, 2) + 0) / 3)) AS cnt            "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and deduct_flg != '*')                                                                   "\
            +"  ) all_cnt                                                                                            "\
            +"  ON list.cnt = all_cnt.cnt                                                                            "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS amt                                                                        "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and deduct_flg != '*')                                                                   "\
            +"  ) ttl                                                                                                "\
            +"  ON 1=1                                                                                               "\
            +" order by amt desc limit 10)                                                                              "\
            +" union all                                                                                                "\
            +" (SELECT '合计',round(sum(max10.amt)/10000,2) as amt,concat(round(100*sum(max10.amt)/ttl.amt,2),'%') pct  "\
            +" FROM (                                                                                                   "\
            +"      SELECT sale_ent_id, SUM(amount) AS amt                                                           "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and deduct_flg != '*')                                                                   "\
            +"      GROUP BY sale_ent_id                                                                             "\
            +"      ORDER BY amt DESC                                                                                "\
            +"      LIMIT 10                                                                                         "\
            +"  ) max10                                                                                              "\
            +"  JOIN (                                                                                               "\
            +"      SELECT sale_ent_id, COUNT(DISTINCT monthsid) AS cnt                                              "\
            +"      FROM (                                                                                           "\
            +"          SELECT DISTINCT sale_ent_id                                                                  "\
            +"              , floor((substring(receipt_date, 6, 2) + 0) / 3) AS monthsid                             "\
            +"          FROM l1_purchase_detail                                                                      "\
            +"          WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"          AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"          AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"          and deduct_flg != '*')                                                               "\
            +"      ) a                                                                                              "\
            +"      GROUP BY sale_ent_id                                                                             "\
            +"  ) list                                                                                               "\
            +"  on max10.sale_ent_id = list.sale_ent_id                                                              "\
            +"  JOIN (                                                                                               "\
            +"      SELECT DISTINCT COUNT(DISTINCT floor((substring(receipt_date, 6, 2) + 0) / 3)) AS cnt            "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and deduct_flg != '*')                                                                   "\
            +"  ) all_cnt                                                                                            "\
            +"  ON list.cnt = all_cnt.cnt                                                                            "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS amt                                                                        "\
            +"      FROM l1_purchase_detail                                                                          "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND deduct_period >= left(date_add('" + dateTm + "', INTERVAL -1 YEAR),7)                         "\
            +"      AND deduct_period < left('" + dateTm + "',7)                                                   "\
            +"      and deduct_flg != '*')                                                                   "\
            +"  ) ttl                                                                                                "\
            +" ON 1 = 1 limit 1);"
    sqlStrList.append(sqlStr3)
    #第四个表格
    sqlStr4 = " (SELECT list.buy_ent_nm, round(max10.ind_value / 10000, 2) AS amt                                        "\
            +"  , concat(round(100 * max10.ind_value / ttl.ind_value, 2), '%') AS pct                                "\
            +" FROM (                                                                                                   "\
            +"  SELECT buy_ent_nm                                                                                    "\
            +"  FROM (                                                                                               "\
            +"      (SELECT buy_ent_nm, SUM(amount) AS ind_value, 1 AS loc                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                "\
            +"      GROUP BY buy_ent_nm                                                                              "\
            +"      ORDER BY ind_value DESC                                                                          "\
            +"      LIMIT 10)                                                                                        "\
            +"      UNION ALL                                                                                        "\
            +"      (SELECT buy_ent_nm, SUM(amount) AS ind_value, 100 AS loc                                                                "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "' "\
            +"            AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -2 YEAR), 7) "\
            +"            AND receipt_mon < LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7))                                                                "\
            +"      GROUP BY buy_ent_nm                                                                              "\
            +"      ORDER BY ind_value DESC                                                                          "\
            +"      LIMIT 10)                                                                                        "\
            +"  ) d                                                                                                  "\
            +"  GROUP BY buy_ent_nm                                                                                  "\
            +"  HAVING COUNT(*) > 1                                                                                  "\
            +" ) list                                                                                                   "\
            +"  JOIN (                                                                                               "\
            +"      SELECT buy_ent_nm, SUM(amount) AS ind_value, 1 AS loc                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                "\
            +"      GROUP BY buy_ent_nm                                                                              "\
            +"      ORDER BY ind_value DESC                                                                          "\
            +"      LIMIT 10                                                                                         "\
            +"  ) max10                                                                                              "\
            +"  ON list.buy_ent_nm = max10.buy_ent_nm                                                                "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS ind_value, 1 AS loc                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                "\
            +"  ) ttl                                                                                                "\
            +"  ON 1 = 1                                                                                             "\
            +"  order by amt desc limit 10)                                                                          "\
            +"  union all                                                                                            "\
            +" (SELECT '合计', round(sum(max10.ind_value) / 10000, 2) AS amt                                            "\
            +"  , concat(round(100 * sum(max10.ind_value) / ttl.ind_value, 2), '%') AS pct                           "\
            +" FROM (                                                                                                   "\
            +"  SELECT buy_ent_nm                                                                                    "\
            +"  FROM (                                                                                               "\
            +"      (SELECT buy_ent_nm, SUM(amount) AS ind_value, 1 AS loc                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                "\
            +"      GROUP BY buy_ent_nm                                                                              "\
            +"      ORDER BY ind_value DESC                                                                          "\
            +"      LIMIT 10)                                                                                        "\
            +"      UNION ALL                                                                                        "\
            +"      (SELECT buy_ent_nm, SUM(amount) AS ind_value, 100 AS loc                                                                "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "' "\
            +"            AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -2 YEAR), 7) "\
            +"            AND receipt_mon < LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7))                                                                "\
            +"      GROUP BY buy_ent_nm                                                                              "\
            +"      ORDER BY ind_value DESC                                                                          "\
            +"      LIMIT 10)                                                                                        "\
            +"  ) d                                                                                                  "\
            +"  GROUP BY buy_ent_nm                                                                                  "\
            +"  HAVING COUNT(*) > 1                                                                                  "\
            +" ) list                                                                                                   "\
            +"  JOIN (                                                                                               "\
            +"      SELECT buy_ent_nm, SUM(amount) AS ind_value, 1 AS loc                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                "\
            +"      GROUP BY buy_ent_nm                                                                              "\
            +"      ORDER BY ind_value DESC                                                                          "\
            +"      LIMIT 10                                                                                         "\
            +"  ) max10                                                                                              "\
            +"  ON list.buy_ent_nm = max10.buy_ent_nm                                                                "\
            +"  JOIN (                                                                                               "\
            +"      SELECT SUM(amount) AS ind_value, 1 AS loc                                                                  "\
            +"      FROM l3_receipt_sum                                                                           "\
            +"      WHERE (ent_tax_id = '" + ent_tax_id + "'                                                                      "\
            +"      AND receipt_mon >= LEFT(date_add('" + dateTm + "', INTERVAL -1 YEAR), 7)                          "\
            +"      AND receipt_mon < LEFT('" + dateTm + "', 7))                                                                "\
            +"  ) ttl                                                                                                "\
            +"  ON 1 = 1 limit 1                                                                                     "\
            +" );                                                                                                    "
    sqlStrList.append(sqlStr4)
    try:
        conn = pymysql.connect(host='127.0.0.1', port=3306, user='root', passwd='yst@2019', db='rdc')
        for sql_name in sqlStrList:
            logging.info(sql_name)
            df = sql.read_sql(sql_name, conn)
            logging.info(df)
            table_list.append(df)
    except Exception as e:
        logging.error(e)
    return table_list


def executeSql_excel(ent_tax_id,dateTm):
    sqlDataList = []
    sqlWordList = []
    sqlStrList = createSqlStr_withParam(ent_tax_id,dateTm)
    try:
        conn = pymysql.connect(host='127.0.0.1', port=3306, user='root', passwd='yst@2019', db='rdc')
        #获取公司名称
        company_sqlStr = "select ent_nm from dim_ent_info where ent_tax_id = '" + ent_tax_id + "'"
        sql_res = sql.read_sql(company_sqlStr,conn)
        print(sql_res)
        company_name = str(sql_res.iloc[0,0])
        for i in range(len(sqlStrList)):
            log_str = "sql num:"+str(i)+":"+sqlStrList[i]
            logging.info(log_str)
            df = sql.read_sql(sqlStrList[i], conn)
            logging.info(df)
            if df.empty:
                sqlDataList.append(0)
            else:
                value = df.iloc[0,1:2].values[0]
                word_value = df.iloc[0,0]
                if value == None:
                    value = 0
                    word_value = 0
                sqlDataList.append(value)
                sqlWordList.append(word_value)
    except ConnectionError as e:
        logging.error(e)
    logging.info(sqlDataList)
    return sqlDataList,sqlWordList,company_name

#指标计算
def calScore(data,length):
    res = np.zeros(length)
    for flag in range(len(data)):
        value = 0
        if flag == 0:
            value = (1-math.exp(-5.5 * (data[flag] + 0.2)))*100
        elif flag == 1 or flag == 2 or flag == 3 or flag == 4 or flag == 5 or flag == 6:
                value = (1 - data[flag]) * 100
        elif flag == 7 or flag == 8 :
            value = data[flag] *10
        elif flag == 9:
            value = (2 - data[flag]) * 50
        elif flag == 10:
            value = (1 - math.exp(-8 * data[flag])) * 100
        elif flag == 11:
            value = (1 - math.exp(-4.5 * data[flag])) * 100
        else :
            value = (1-math.exp(-11 * data[flag]))*100
        #确保所有的值位于1-100，超过的值需要截断
        if value < 0:
            value = 0
        elif value > 100:
            value = 100
        else:
            value = round(value,2)
        res[flag] = value
    return res

#获取Excel中sql执行和计算结果
def enterExcel(ent_tax_id,dateTm):
 #从数据库中获取数据
    value,value_doc,company_name = executeSql_excel(ent_tax_id,dateTm)
    df = pd.read_excel("D:/uploadTempFiles/python/reportTemp/quantitativeAnalysisTemplate.xls",skiprows=1)
    # print(df.iloc[:,3:4])
    #从模板中获取指标权重
    resDataFrame = df.iloc[:,3:4]
    # 构建excel中需要的数组
    score_list = [np.nan for i in range(df.shape[0])]
    value_list = [np.nan for i in range(df.shape[0])]
    # 分数计算
    score_data = calScore(value,df.shape[0])
    # print(score_data)
    #构建excel中的的值、分数和乘积列
    data_flag = 0
    for i in range(len(score_data)):
        if i == 1 or i == 8 or i == 12 or i == 16 or i == 17 or i== 18:
            continue
        else:
            score_list[i] = score_data[data_flag]
            value_list[i] = value[data_flag]
            data_flag += 1
    resDataFrame["value"] = value_list
    resDataFrame["score"] = score_list
    resDataFrame["res"] = resDataFrame.iloc[:,[0,2]].fillna(0).prod(axis=1)
    #求乘积的汇总值
    total = resDataFrame["res"].sum()
    resDataFrame.iloc[1, 3] = resDataFrame.iloc[0, 3]
    resDataFrame.iloc[8, 3] = resDataFrame.iloc[2:8, 3].sum()
    resDataFrame.iloc[12, 3] = resDataFrame.iloc[9:12, 3].sum()
    resDataFrame.iloc[16, 3] = resDataFrame.iloc[13:16, 3].sum()
    resDataFrame.iloc[17, 3] = total
    # print(resDataFrame)
    
    #excel模板填充
    book = xlrd.open_workbook("D:/uploadTempFiles/python/reportTemp/quantitativeAnalysisTemplate.xls", formatting_info=True)
    new_book = copy(book)
    sheet = new_book.get_sheet(0)
    #第一行需要加上公司名称
    title_name = "                                                          "+company_name+"定量分析"
    sheet.write(0,0,title_name)
    for rowNum in range(df.shape[0]):
        for colNum in range(4):
            if np.isnan(resDataFrame.iloc[rowNum, colNum]) == False:
                print(resDataFrame.iloc[rowNum, colNum])
                # 给特定列加上%
                if colNum == 0:
                    value = str(round(resDataFrame.iloc[rowNum, colNum] * 100, 2)) + "%"
                elif colNum == 1:
                    if rowNum == 9  or rowNum == 10 or rowNum == 11:
                        value = str(round(resDataFrame.iloc[rowNum, colNum]))
                    else:
                        value = str(round(resDataFrame.iloc[rowNum, colNum] * 100, 2)) + "%"
                else:
                    value = str(round(resDataFrame.iloc[rowNum, colNum], 2))
                sheet.write(rowNum + 2, colNum + 3, value) # 写入excel，第一个值是行，第二个值是列
    company_name = company_name.replace("?","")
    file_name = "D:/uploadTempFiles/python/reportTemp/" +company_name+"_"+date.today().strftime("%Y%m%d")+ "_定量评分"+ ".xls"
    new_book.save(file_name)  # 保存新的excel，保存excel必须使用后缀名是.xls的，不是能是.xlsx的
    data = xlrd.open_workbook(file_name)
    table = data.sheet_by_name("Sheet1")
    nrows = 20
    result = []
    result.append(ent_tax_id)
    result.append(company_name)
    for j in [2,4,5,6,7,8,9,11,12,13,15,16,17]:
        rows = table.row_values(j)
        result.append(rows[5])
    result.append(table.row_values(19)[6])
    # with open(r"D:\uploadTempFiles\python\logFile\entScore.txt","a") as f:
        # f.write(",".join(result)+"\n")
    try:
        conn_score = pymysql.connect(host='127.0.0.1', port=3306, user='root', passwd='yst@2019',db='rdc')
        cur_score = conn_score.cursor()
        insert_sql = "insert into rdc.res_ent_score values(current_timestamp(),'"+ent_tax_id+"','"+company_name+"',"+",".join(result[2:])+",'processing');"
        cur_score.execute(insert_sql)
        conn_score.commit()
        cur_score.close()
        conn_score.close()
    except Exception as e:
        print(insert_sql+"\n"+str(e)+"\n")
        with open(r"D:\uploadTempFiles\python\logFile\log.txt","a") as f:
            f.write(company_name+'\n'+str(e)+'\n')
    try:
        shutil.copyfile(file_name,"D:/report/ori/daiqian/" +company_name+"_"+date.today().strftime("%Y%m%d")+ "_定量评分.xls")
        # excelToPDF("D:/report/ori/daiqian/" +company_name+"_"+date.today().strftime("%Y%m%d")+ "_定量评分.xls")
    except Exception as e:
        with open(r"D:\uploadTempFiles\python\logFile\log.txt","a") as f:
            f.write(company_name+'\n'+str(e)+'\n')

    return resDataFrame,value_doc

def executeSql_doc(ent_tax_id,dateTm):
    sqlDataList = ["#" for i in range(70)]
    sqlStrList = createSqlStr_for_word(ent_tax_id,dateTm)
    try:
        conn = pymysql.connect(host='127.0.0.1', port=3306, user='root', passwd='yst@2019', db='rdc')
        for i in range(len(sqlStrList)):
            sqlNum = i + 1
            logging.info(sqlNum)
            logging.info(sqlStrList[i])
            df = sql.read_sql(sqlStrList[i], conn)
            logging.info(df)
            if i == 0:
                if df.empty == False:
                    sqlDataList[0] = df.iloc[0, 0]
                    sqlDataList[1] = df.iloc[0, 1]
                    sqlDataList[53] = df.iloc[0, 2]
            elif i == 1:
                if df.empty == False:
                    sqlDataList[9] = df.iloc[0, 1]
                    if(int(sqlDataList[9])<10):
                        sqlDataList[11]=""
                    else:
                        sqlDataList[11]="其中前十大"
                    sqlDataList[10] = df.iloc[0, 2]
            elif i == 2:
                if df.empty == False:
                    sqlDataList[5] = df.iloc[0, 2]
                    sqlDataList[16] = df.iloc[0, 1]
                    if(int(sqlDataList[16])<10):
                        sqlDataList[66]=""
                    else:
                        sqlDataList[66]="其中前十大"
                    sqlDataList[17] = df.iloc[0, 2]
            elif i == 3:
                if df.empty == False:
                    sqlDataList[23] = df.iloc[0, 2]
            elif i == 4:
                if df.empty == False:
                    sqlDataList[32] = str(round(df.iloc[0, 1]*100,2)) + "%"
            elif i == 5:
                if df.empty == False:
                    sqlDataList[37] = df.iloc[0, 0]
            elif i == 6:
                if df.empty == False:
                    sqlDataList[42] = df.iloc[0, 0]
            elif i == 7:
                if df.empty == False:
                    sqlDataList[49] = df.iloc[0, 1]
                    sqlDataList[51] = df.iloc[0, 0]
            elif i == 8:
                if df.empty == False:
                    sqlDataList[56] = df.iloc[0, 1]
                    sqlDataList[58] = df.iloc[0, 0]
            else:
                if df.empty == False:
                    sqlDataList[62] = df.iloc[0, 1]
                    sqlDataList[64] = df.iloc[0, 0]
        conn.close()
        # 时间指标
        stamp = datetime.strptime(dateTm, '%Y-%m-%d')
        last_one_year = stamp - pd.tseries.offsets.DateOffset(months=12)
        last_two_year = stamp - pd.tseries.offsets.DateOffset(months=24)
        # 时间格式化,以点为标准
        last_one_year = last_one_year.strftime('%Y-%m-%d')
        last_two_year = last_two_year.strftime('%Y-%m-%d')
        sqlDataList[2] = str(last_two_year) + "-" + dateTm
        sqlDataList[22] = str(last_one_year) + "-" + dateTm
        sqlDataList[27] = str(last_one_year)[0:4]
        sqlDataList[31] = str(last_one_year)[0:4]
        sqlDataList[48] = str(last_two_year) + "-" + str(last_one_year)
        sqlDataList[50] = str(last_one_year) + "-" + dateTm
        sqlDataList[55] = str(last_two_year) + "-" + str(last_one_year)
        sqlDataList[57] = str(last_one_year) + "-" + dateTm
        sqlDataList[61] = str(last_two_year) + "-" + str(last_one_year)
        sqlDataList[63] = str(last_one_year) + "-" + dateTm
        #模板中的固定值
        # sqlDataList[11] = "其中前十大"
        #Excel模板中的部分指标
        excel_df,excel_df_src = enterExcel(ent_tax_id,dateTm)
        # 行业平均销售额增长率
        sqlDataList[3] = str(round(excel_df.iloc[0, 1]*100,2)) + "%"
        sqlDataList[4] = excel_df.iloc[0, 3]
        # 销售额本地行业地位
        sqlDataList[6] = int(excel_df_src[1])
        sqlDataList[7] = str(round(excel_df.iloc[2, 1]*100,2)) + "%"
        sqlDataList[8] = excel_df.iloc[2, 3]
        # 上游交易对手合计数量本地行业地位
        sqlDataList[13] = int(excel_df_src[2])
        sqlDataList[14] = str(round(excel_df.iloc[3, 1]*100,2)) + "%"
        sqlDataList[15] = excel_df.iloc[3, 3]
        # 下游交易对手合计数量本地行业地位
        sqlDataList[19] = int(excel_df_src[3])
        sqlDataList[20] = str(round(excel_df.iloc[4, 1]*100,2)) + "%"
        sqlDataList[21] = excel_df.iloc[4, 3]
        # 近一年增值税纳税额本地行业地位
        # 增值税税负率本地行业地位
        if(sqlDataList[23])>0:
            sqlDataList[24] = "，在本地该行业排名第" + str(int(excel_df_src[4])) +"位"
            sqlDataList[25] = "，处于该行业前" + str(round(excel_df.iloc[5, 1]*100,2)) + "%"
            sqlDataList[28] = "在本地该行业排名第" + str(int(excel_df_src[5])) + "位"
            sqlDataList[29] = "，处于该行业前" + str(round(excel_df.iloc[6, 1]*100,2)) + "%"
        else:
            sqlDataList[24] = ""
            sqlDataList[25] = ""
            sqlDataList[28] = ""
            sqlDataList[29] = ""
        sqlDataList[26] = excel_df.iloc[5, 3]
        sqlDataList[30] = excel_df.iloc[6, 3]
        # 毛利率本地行业地位
        sqlDataList[33] = int(excel_df_src[6])
        sqlDataList[34] = str(round(excel_df.iloc[7, 1]*100,2)) + "%"
        sqlDataList[35] = excel_df.iloc[7, 3]
        # 近一年最大十家稳定上游供应商数量
        sqlDataList[36] = int(excel_df.iloc[9, 1])
        sqlDataList[38] = excel_df.iloc[9, 3]
        sqlDataList[39] = "这些"
        # 近两年每年最大十家稳定下游交易企业数量
        sqlDataList[41] = int(excel_df.iloc[10, 1])
        sqlDataList[43] = excel_df.iloc[10, 3]
        sqlDataList[44] = "这些"
        # 断档
        sqlDataList[46] = excel_df.iloc[11, 3]
        if(excel_df.iloc[11,1]>0):
            sqlDataList[67] = "有"+str(int(excel_df.iloc[11,1]))+"次"
            pattern = re.compile(r'-[0]{0,1}')
            sqlDataList[68] = "分别是"+ re.sub(pattern, "年",excel_df_src[9]).replace("，","月，") +"月。"
        else:
            sqlDataList[67] = "无"
            sqlDataList[68] = ""
        # 销售额增长率
        sqlDataList[47] = str(round(excel_df.iloc[13, 1]*100,2)) + "%"
        sqlDataList[52] = excel_df.iloc[13, 3]
        # 主营商品销售额增长率
        sqlDataList[54] = str(round(excel_df.iloc[14, 1]*100,2)) + "%"
        sqlDataList[59] = excel_df.iloc[14, 3]
        # 增值税纳税额增长率
        sqlDataList[60] = str(round(excel_df.iloc[15, 1]*100,2)) + "%"
        sqlDataList[65] = excel_df.iloc[15, 3]

        ##表格数据赋值
        sqlDataList[12] = "@1"
        sqlDataList[18] = "@2"
        sqlDataList[40] = "@3"
        sqlDataList[45] = "@4"

        # 企业评分
        sqlDataList[69] = excel_df.iloc[17, 3]
    except ConnectionError as e:
        logging.error(e)
    logging.info(sqlDataList)
    return sqlDataList

def test_main(ent_tax_id, dateTm):
    d = Document('D:/uploadTempFiles/python/reportTemp/doc_input.docx')
    table_list = get_table_data(ent_tax_id, dateTm)
    doc = Document()
    remove_re = re.compile(u'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]')
    # 表格插入
    for i in range(len(d.paragraphs)):
        # 标题
        if i == 0:
            p = doc.add_paragraph(d.paragraphs[i].text)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # run = doc.add_paragraph(d.paragraphs[i].text).add_run()
        elif str(d.paragraphs[i].text).find("、") >= 0:
            p = doc.add_paragraph()
            p.add_run(d.paragraphs[i].text).bold = True
        elif d.paragraphs[i].text == "@1":
            df = table_list[0]
            table_rownum = df.shape[0] + 1
            table = doc.add_table(table_rownum, 4, style="Table Grid")
            # 表头
            heading_cells = table.rows[0].cells
            heading_cells[0].text = '序号'
            heading_cells[1].text = '上游供应商名称'
            heading_cells[2].text = '进项金额（万元）'
            heading_cells[3].text = '占比'
            if table_rownum > 2:
                for rownum in range(df.shape[0]):
                    # print(rownum + 1)
                    heading_cells = table.rows[rownum + 1].cells
                    heading_cells[0].text = str(rownum + 1)
                    heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                    heading_cells[2].text = str(df.iloc[rownum, 1])
                    heading_cells[3].text = df.iloc[rownum, 2]
        elif d.paragraphs[i].text == "@2":
            df = table_list[1]
            table_rownum = df.shape[0] + 1
            table = doc.add_table(table_rownum, 4, style="Table Grid")
            # 表头
            heading_cells = table.rows[0].cells
            heading_cells[0].text = '序号'
            heading_cells[1].text = '下游交易对手名称'
            heading_cells[2].text = '交易金额（万元）'
            heading_cells[3].text = '占比'
            if table_rownum > 2:
                for rownum in range(df.shape[0]):
                    # print(rownum + 1)
                    heading_cells = table.rows[rownum + 1].cells
                    heading_cells[0].text = str(rownum + 1)
                    heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                    heading_cells[2].text = str(df.iloc[rownum, 1])
                    heading_cells[3].text = df.iloc[rownum, 2]
        elif d.paragraphs[i].text == "@3":
            df = table_list[2]
            table_rownum = df.shape[0] + 1
            table = doc.add_table(table_rownum, 4, style="Table Grid")
            # 表头
            heading_cells = table.rows[0].cells
            heading_cells[0].text = '序号'
            heading_cells[1].text = '上游供应商名称'
            heading_cells[2].text = '进项金额（万元）'
            heading_cells[3].text = '占比'
            if table_rownum > 2:
                for rownum in range(df.shape[0]):
                    # print(rownum + 1)
                    heading_cells = table.rows[rownum + 1].cells
                    heading_cells[0].text = str(rownum + 1)
                    heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                    heading_cells[2].text = str(df.iloc[rownum, 1])
                    heading_cells[3].text = df.iloc[rownum, 2]
        elif d.paragraphs[i].text == "@4":
            df = table_list[3]
            table_rownum = df.shape[0] + 1
            table = doc.add_table(table_rownum, 4, style="Table Grid")
            # 表头
            heading_cells = table.rows[0].cells
            heading_cells[0].text = '序号'
            heading_cells[1].text = '下游交易对手名称'
            heading_cells[2].text = '交易金额（万元）'
            heading_cells[3].text = '占比'
            if table_rownum > 2:
                for rownum in range(df.shape[0]):
                    # print(rownum + 1)
                    heading_cells = table.rows[rownum + 1].cells
                    heading_cells[0].text = str(rownum + 1)
                    heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                    heading_cells[2].text = str(df.iloc[rownum, 1])
                    heading_cells[3].text = df.iloc[rownum, 2]
        else:
            # print(d.paragraphs[i].text)
            doc.add_paragraph(d.paragraphs[i].text)
    doc.save('D:/uploadTempFiles/python/reportTemp/doc_output.docx')

def get_ent_nm(ent_tax_id):
    url = 'http://www.qichacha.com/search?key='+ent_tax_id
    header = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.62 Safari/537.36',
        'Cookie': 'UM_distinctid=16a0bf964b2af-093c01fdf13f0d-38395e0b-140000-16a0bf964b3195; zg_did=%7B%22did%22%3A%20%2216a0bf9650b375-04b14a91ff1a74-38395e0b-140000-16a0bf9650c1a4%22%7D; _uab_collina=155497905533903516796765; QCCSESSID=v155bo7j560naf2ok7eg3rc6u6; CNZZDATA1254842228=469224891-1554975454-https%253A%252F%252Fwww.baidu.com%252F%7C1561369279; hasShow=1; Hm_lvt_3456bee468c83cc63fb5147f119f1075=1558950734,1560172590,1561370779; acw_tc=6f7b369d15613708179653436e61ef65c0d32b44f79b91a08a5be69371; zg_de1d1a35bfa24ce29bbf2c7eb17e6c4f=%7B%22sid%22%3A%201561370777682%2C%22updated%22%3A%201561370804892%2C%22info%22%3A%201561370777690%2C%22superProperty%22%3A%20%22%7B%7D%22%2C%22platform%22%3A%20%22%7B%7D%22%2C%22utm%22%3A%20%22%7B%7D%22%2C%22referrerDomain%22%3A%20%22%22%2C%22cuid%22%3A%20%223a1df6b296369b4b5fef33438e007bf2%22%7D; Hm_lpvt_3456bee468c83cc63fb5147f119f1075=1561370805'
    }
    try:
        html = requests.get(url, headers=header, )
        page = etree.HTML(html.text)
    
        ent_nm = page.xpath('//*[@id="search-result"]/tr/td[3]/a/text()')[0]
    except Exception as e:
        ent_nm = ent_tax_id
    return ent_nm

def enterWord(ent_tax_id, dateTm):
    value_list = executeSql_doc(ent_tax_id, dateTm)
    table_list = get_table_data(ent_tax_id, dateTm)
    remove_re = re.compile(u'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]')
    out_table_list = []
    # 对插入数据进行预处理，去除None值
    for i in range(len(value_list)):
        if value_list[i] == None:
            value_list[i] = 0
        elif isinstance(value_list[i],float):
            value_list[i] = round(value_list[i],2)
    #控制表格是否输出,0代表不输出
    
    buy_ent_nm_dict = {}
    for table_num in range(len(table_list)):
        df = table_list[table_num]
        # 对进项企业，id改为名称
        if(table_num == 0):
            for j in range(len(df)-1):
                buy_ent_nm_dict[df.iloc[j,0]] = get_ent_nm(df.iloc[j,0])
                df.iloc[j,0] = buy_ent_nm_dict[df.iloc[j,0]]
        if(table_num == 2):
            for j in range(len(df)-1):
                df.iloc[j,0] = buy_ent_nm_dict[df.iloc[j,0]]
        if  df.shape[0] < 2:
            out_table_list.append(0)
            if table_num == 0:
                value_list[12] = ""
            elif table_num == 1:
                value_list[18] = ""
            elif table_num == 2:
                value_list[40] = ""
            else:
                value_list[45] = ""
        else:
            out_table_list.append(1)
    # print(out_table_list)
    try:
        file_path = "D:/uploadTempFiles/python/reportTemp/" +value_list[0] +"_"+ date.today().strftime("%Y%m%d")+"_定量分析.doc"
        d = Document('D:/uploadTempFiles/python/reportTemp/quantitativeAnalysisTemplate.docx')
        
        # 写企业信息
        conn_yst = pymysql.connect(host='127.0.0.1', port=3306, user='root', passwd='yst@2019', db='yst')
        company_sqlStr = "select ent_business_staff, ent_business_phone, ent_real_address,loan_amt, register_name, register_phone\
        from loan_info loan left join register_auth_info reg on loan.ent_tax_id = reg.ent_tax_id \
        where loan.ent_tax_id = '"+ent_tax_id+"' and (apply_status='2' or apply_status is null);"
        sql_res = sql.read_sql(company_sqlStr,conn_yst)
        print(sql_res)
        company_info = "企业法人："+str(sql_res.iloc[0,4])+"，法人联系方式："+str(sql_res.iloc[0,5])+\
        "。\n业务联系人："+str(sql_res.iloc[0,0])+"，业务联系方式："+str(sql_res.iloc[0,1])+"。\n企业地址："+str(sql_res.iloc[0,2])+"。意愿贷款金额："+str(sql_res.iloc[0,3])+"万元。"
        conn_yst.close()
        
        # 遍历每段，在每段中执行替换动作
        for para in d.paragraphs:
            # print(para.text)
            for i in range(len(value_list) - 1, -1, -1):
                old_value = "#" + str(i)
                para.text = para.text.replace(old_value, str(value_list[i]))
            para.text = para.text.replace("#enterprise information#",company_info)
        doc = Document()
        table_number = 0
        # 表格插入
        for i in range(len(d.paragraphs)):
            # 标题
            if i == 0:
                p = doc.add_paragraph()
                p.add_run(d.paragraphs[i].text).bold = True
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # run = doc.add_paragraph(d.paragraphs[i].text).add_run()
            elif str(d.paragraphs[i].text).find("、") >= 0 :
                p = doc.add_paragraph()
                p.add_run(d.paragraphs[i].text).bold = True
            #控制表格上方的是否输出
            elif str(d.paragraphs[i].text).find("情况如下") >= 0 :
                    logging.info("情况如下:"+str(table_number))
                    if out_table_list[table_number] == 1:
                        doc.add_paragraph(d.paragraphs[i].text)
                    table_number += 1

            elif d.paragraphs[i].text == "@1":
                df = table_list[0]
                table_rownum = df.shape[0] + 1
                table = doc.add_table(table_rownum, 4, style="Table Grid")
                # 表头
                heading_cells = table.rows[0].cells
                heading_cells[0].text = '序号'
                heading_cells[1].text = '上游供应商名称'
                heading_cells[2].text = '进项金额（万元）'
                heading_cells[3].text = '占比'
                if table_rownum > 2:
                    for rownum in range(df.shape[0]):
                        # print(rownum + 1)
                        heading_cells = table.rows[rownum + 1].cells
                        heading_cells[0].text = str(rownum + 1)
                        heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                        heading_cells[2].text = str(df.iloc[rownum, 1])
                        heading_cells[3].text = df.iloc[rownum, 2]
            elif d.paragraphs[i].text == "@2":
                df = table_list[1]
                table_rownum = df.shape[0] + 1
                table = doc.add_table(table_rownum, 4, style="Table Grid")
                # 表头
                heading_cells = table.rows[0].cells
                heading_cells[0].text = '序号'
                heading_cells[1].text = '下游交易对手名称'
                heading_cells[2].text = '交易金额（万元）'
                heading_cells[3].text = '占比'
                if table_rownum > 2:
                    for rownum in range(df.shape[0]):
                        # print(rownum + 1)
                        heading_cells = table.rows[rownum + 1].cells
                        heading_cells[0].text = str(rownum + 1)
                        heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                        heading_cells[2].text = str(df.iloc[rownum, 1])
                        heading_cells[3].text = df.iloc[rownum, 2]
            elif d.paragraphs[i].text == "@3":
                df = table_list[2]
                table_rownum = df.shape[0] + 1
                table = doc.add_table(table_rownum, 4, style="Table Grid")
                # 表头
                heading_cells = table.rows[0].cells
                heading_cells[0].text = '序号'
                heading_cells[1].text = '上游供应商名称'
                heading_cells[2].text = '进项金额（万元）'
                heading_cells[3].text = '占比'
                if table_rownum > 2:
                    for rownum in range(df.shape[0]):
                        # print(rownum + 1)
                        heading_cells = table.rows[rownum + 1].cells
                        heading_cells[0].text = str(rownum + 1)
                        heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                        heading_cells[2].text = str(df.iloc[rownum, 1])
                        heading_cells[3].text = df.iloc[rownum, 2]
            elif d.paragraphs[i].text == "@4":
                df = table_list[3]
                table_rownum = df.shape[0] + 1
                table = doc.add_table(table_rownum, 4, style="Table Grid")
                # 表头
                heading_cells = table.rows[0].cells
                heading_cells[0].text = '序号'
                heading_cells[1].text = '下游交易对手名称'
                heading_cells[2].text = '交易金额（万元）'
                heading_cells[3].text = '占比'
                if table_rownum > 2:
                    for rownum in range(df.shape[0]):
                        # print(rownum + 1)
                        heading_cells = table.rows[rownum + 1].cells
                        heading_cells[0].text = str(rownum + 1)
                        heading_cells[1].text = remove_re.sub("",df.iloc[rownum, 0])
                        heading_cells[2].text = str(df.iloc[rownum, 1])
                        heading_cells[3].text = df.iloc[rownum, 2]
            else:
                # print(d.paragraphs[i].text)
                doc.add_paragraph(d.paragraphs[i].text)
        doc.save(file_path)
        shutil.copyfile(file_path,"D:/report/ori/daiqian/" +value_list[0] +"_"+ date.today().strftime("%Y%m%d")+"_定量分析.doc")
        # docToPDF("D:/report/ori/daiqian/" +value_list[0] +"_"+ date.today().strftime("%Y%m%d")+"_定量分析.doc")
        return value_list[0],value_list[69]
    except Exception as e:
        logging.error(e)
        print(e)
        return 'error',-1


def toReport_main(ent_tax_id, end_date):
    time_start = time.time()
    ent_nm,score = enterWord(ent_tax_id,end_date)
    time_end=time.time()
    try:
        conn_time = pymysql.connect(host='127.0.0.1', port=3306, user='root', passwd='yst@2019',db='rdc')
        cur_time = conn_time.cursor()
        update_sql = "update rdc.res_ent_score set time_cost = '"+str(round(time_end-time_start,2))+"' where ent_tax_id = '"+ent_tax_id+"' and cost_time = 'processing');"
        cur_time.execute(update_sql)
        conn_time.commit()
        cur_time.close()
        conn_time.close()
        print("---- totally cost "+str(round(time_end-time_start,2))+" seconds ----")
    except Exception as e:
        print(update_sql+"\n"+str(e)+"\n")
        with open(r"D:\uploadTempFiles\python\logFile\log.txt","a") as f:
            f.write(update_sql+'\n'+str(e)+'\n')

    return ent_nm,score

def docToPDF(file_path):
    pdf_path = file_path.replace(".doc",".pdf")
    stat = os.system('taskkill /F /im wps.exe')
    o = win32com.client.Dispatch("Kwps.Application")
    o.Visible=False
    doc = o.Documents.Open(file_path)
    doc.ExportAsFixedFormat(pdf_path,17)
    o.Quit()

def excelToPDF(file_path):
    pdf_path = file_path.replace(".xls",".pdf")
    stat = os.system('taskkill /F /im wps.exe')
    xlApp = win32com.client.Dispatch("Ket.Application")
    excel = xlApp.Workbooks.Open(file_path)
    excel.ExportAsFixedFormat(0, pdf_path)
    excel.Close()
    xlApp.Quit()

if __name__ == "__main__":
    toReport_main(sys.argv[1],sys.argv[2])

